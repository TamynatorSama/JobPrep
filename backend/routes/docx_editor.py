"""In-place .docx section editor.

Given the raw bytes of a master resume .docx and a `sections` dict produced
by the LLM, walk the source document, replace each known section's body with
the new content, and preserve original styling (fonts, bold runs, paragraph
indentation, bullet numbering, colors).

The editor is intentionally conservative: if the source document uses tables
for layout or its section headers cannot be confidently identified, the
caller should fall back to rebuilding the resume from scratch.

Public API:
    apply_edits_to_docx(source_bytes, sections) -> bytes
    sections_to_text(sections) -> str   # markdown text used by from-scratch
                                          # renderer and chat preview
"""
from __future__ import annotations

import io
from copy import deepcopy
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ── section detection ───────────────────────────────────────────────────────

SECTION_ALIASES: dict[str, list[str]] = {
    "summary":        ["summary", "professional summary", "profile", "about", "about me", "objective"],
    "skills":         ["skills", "technical skills", "expertise", "tech stack", "skill set"],
    "experience":     ["experience", "professional experience", "work experience", "employment", "employment history", "work history"],
    "projects":       ["projects", "key projects", "selected projects", "personal projects", "notable projects"],
    "education":      ["education", "academic background", "academic qualifications"],
    "certifications": ["certifications", "certificates", "licenses", "licences"],
    "achievements":   ["achievements", "awards", "honors", "honours", "accomplishments"],
}


def _normalize(text: str) -> str:
    return text.strip().lower().rstrip(":").strip()


def _looks_like_header(p) -> bool:
    """Heuristic: short text, bold OR Word Heading style."""
    txt = p.text.strip()
    if not txt or len(txt) > 60:
        return False
    style_name = (p.style.name if p.style else "").lower()
    if "heading" in style_name or style_name.startswith("title"):
        return True
    return any(r.bold for r in p.runs if r.bold is not None)


def _is_bullet(p) -> bool:
    pPr = p._element.find(qn("w:pPr"))
    if pPr is None:
        return False
    return pPr.find(qn("w:numPr")) is not None


def _is_subhead(p) -> bool:
    """A non-bullet paragraph with bold text — used as role/project header."""
    if _is_bullet(p):
        return False
    if not p.text.strip():
        return False
    return any(r.bold for r in p.runs if r.bold)


def _identify_sections(paras) -> dict[str, tuple[int, int]]:
    """Map canonical section name -> (header_idx, end_idx_exclusive)."""
    headers: list[tuple[str, int]] = []
    for i, p in enumerate(paras):
        norm = _normalize(p.text)
        if not norm:
            continue
        for canon, aliases in SECTION_ALIASES.items():
            if norm in aliases and _looks_like_header(p):
                headers.append((canon, i))
                break
    out: dict[str, tuple[int, int]] = {}
    for k, (canon, idx) in enumerate(headers):
        next_idx = headers[k + 1][1] if k + 1 < len(headers) else len(paras)
        out[canon] = (idx, next_idx)
    return out


# ── XML helpers ─────────────────────────────────────────────────────────────

def _first_run_rPr(para) -> Optional[object]:
    p_elem = para._element
    first_r = p_elem.find(qn("w:r"))
    if first_r is None:
        return None
    rPr = first_r.find(qn("w:rPr"))
    return deepcopy(rPr) if rPr is not None else None


def _clone_paragraph_shell(template_para) -> object:
    """Return a deep-copied <w:p> element with all <w:r> runs stripped.
    Paragraph-level properties (pPr: alignment, numbering, indent) are kept.
    """
    new_p = deepcopy(template_para._element)
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    return new_p


def _build_run(text: str, rPr_template: Optional[object], force_bold: bool = False) -> object:
    """Build a <w:r> with optional cloned rPr + text."""
    r = OxmlElement("w:r")
    if rPr_template is not None:
        rPr = deepcopy(rPr_template)
    else:
        rPr = OxmlElement("w:rPr") if force_bold else None
    if force_bold:
        if rPr is None:
            rPr = OxmlElement("w:rPr")
        # Replace any existing <w:b> with <w:b w:val="1"/>
        for old in rPr.findall(qn("w:b")):
            rPr.remove(old)
        b = OxmlElement("w:b")
        rPr.append(b)
    if rPr is not None:
        r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


def _set_paragraph_text(para, new_text: str) -> None:
    """Replace text inside `para` while keeping the first run's rPr."""
    p_elem = para._element
    rPr = _first_run_rPr(para)
    for r in p_elem.findall(qn("w:r")):
        p_elem.remove(r)
    p_elem.append(_build_run(new_text, rPr))


def _insert_clone(cursor, template_para, text: str, *, force_bold: bool = False):
    """Insert a paragraph after `cursor` cloning `template_para`'s style with
    `text`. Returns the new <w:p> element to chain further insertions.
    """
    if template_para is None:
        new_p = OxmlElement("w:p")
        new_p.append(_build_run(text, None, force_bold=force_bold))
        cursor.addnext(new_p)
        return new_p
    new_p = _clone_paragraph_shell(template_para)
    rPr = _first_run_rPr(template_para)
    new_p.append(_build_run(text, rPr, force_bold=force_bold))
    cursor.addnext(new_p)
    return new_p


def _insert_two_run(cursor, template_para, bold_prefix: str, plain_suffix: str):
    """Insert a paragraph with `bold_prefix` in bold + `plain_suffix` plain."""
    if template_para is None:
        new_p = OxmlElement("w:p")
        new_p.append(_build_run(bold_prefix, None, force_bold=True))
        new_p.append(_build_run(plain_suffix, None))
        cursor.addnext(new_p)
        return new_p
    new_p = _clone_paragraph_shell(template_para)
    base_rPr = _first_run_rPr(template_para)
    new_p.append(_build_run(bold_prefix, base_rPr, force_bold=True))
    new_p.append(_build_run(plain_suffix, base_rPr))
    cursor.addnext(new_p)
    return new_p


# ── section emitters ────────────────────────────────────────────────────────

def _capture_templates(body_paras):
    """Pull representative paragraphs (subhead, bullet, plain) from a section
    body to reuse as styling templates.
    """
    bullet_t = next((p for p in body_paras if _is_bullet(p)), None)
    subhead_t = next((p for p in body_paras if _is_subhead(p) and not _is_bullet(p)), None)
    plain_t = next(
        (p for p in body_paras if not _is_bullet(p) and not _is_subhead(p) and p.text.strip()),
        None,
    )
    fallback = body_paras[0] if body_paras else None
    return bullet_t, subhead_t, plain_t, fallback


def _emit_summary(cursor, data, bullet_t, subhead_t, plain_t, fb):
    text = str(data or "").strip()
    if not text:
        return cursor
    return _insert_clone(cursor, plain_t or fb, text)


def _emit_skills(cursor, data, bullet_t, subhead_t, plain_t, fb):
    for cat in data or []:
        category = str(cat.get("category", "") or "").strip()
        items = cat.get("items", []) or []
        items = [str(x).strip() for x in items if str(x).strip()]
        if not items:
            continue
        prefix = f"{category}: " if category else ""
        suffix = ", ".join(items)
        if prefix:
            cursor = _insert_two_run(cursor, plain_t or fb, prefix, suffix)
        else:
            cursor = _insert_clone(cursor, plain_t or fb, suffix)
    return cursor


def _format_exp_head(item: dict) -> str:
    role = str(item.get("role", "") or "").strip()
    company = str(item.get("company", "") or "").strip()
    location = str(item.get("location", "") or "").strip()
    start = str(item.get("start", "") or "").strip()
    end = str(item.get("end", "") or "").strip()
    head = role
    if company:
        head = f"{head} — {company}" if head else company
    if location:
        head = f"{head}, {location}"
    if start or end:
        dates = f"{start} – {end}".strip(" –")
        head = f"{head} · {dates}" if head else dates
    return head


def _format_project_head(item: dict) -> str:
    name = str(item.get("name", "") or "").strip()
    subtitle = str(item.get("subtitle", "") or "").strip()
    date = str(item.get("date", "") or "").strip()
    head = name
    if subtitle:
        head = f"{head} — {subtitle}" if head else subtitle
    if date:
        head = f"{head} · {date}" if head else date
    return head


def _emit_roles(cursor, data, bullet_t, subhead_t, plain_t, fb, head_fn):
    for item in data or []:
        head = head_fn(item)
        if head:
            cursor = _insert_clone(cursor, subhead_t or plain_t or fb, head, force_bold=(subhead_t is None))
        for bullet in item.get("bullets", []) or []:
            text = str(bullet).strip()
            if not text:
                continue
            cursor = _insert_clone(cursor, bullet_t or plain_t or fb, text)
    return cursor


def _emit_education(cursor, data, bullet_t, subhead_t, plain_t, fb):
    for item in data or []:
        inst = str(item.get("institution", "") or "").strip()
        loc = str(item.get("location", "") or "").strip()
        head = f"{inst} — {loc}" if inst and loc else inst or loc
        if head:
            cursor = _insert_clone(cursor, subhead_t or plain_t or fb, head, force_bold=(subhead_t is None))
        body_bits = []
        deg = str(item.get("degree", "") or "").strip()
        if deg:
            body_bits.append(deg)
        dates = str(item.get("dates", "") or "").strip()
        if dates:
            body_bits.append(dates)
        if body_bits:
            cursor = _insert_clone(cursor, plain_t or fb, " · ".join(body_bits))
    return cursor


def _emit_list_items(cursor, data, bullet_t, subhead_t, plain_t, fb):
    for item in data or []:
        text = str(item).strip()
        if not text:
            continue
        cursor = _insert_clone(cursor, bullet_t or plain_t or fb, text)
    return cursor


SECTION_EMITTERS = {
    "summary":        _emit_summary,
    "skills":         _emit_skills,
    "experience":     lambda c, d, b, s, p, f: _emit_roles(c, d, b, s, p, f, _format_exp_head),
    "projects":       lambda c, d, b, s, p, f: _emit_roles(c, d, b, s, p, f, _format_project_head),
    "education":      _emit_education,
    "certifications": _emit_list_items,
    "achievements":   _emit_list_items,
}


# ── public API ──────────────────────────────────────────────────────────────

def apply_edits_to_docx(source_bytes: bytes, sections: dict) -> bytes:
    """Edit a master resume .docx in place using `sections` data.

    Strategy:
      1. Open source with python-docx.
      2. Bail out (raise ValueError) if the layout is table-based — caller
         falls back to the from-scratch renderer.
      3. Replace the first two non-empty paragraphs with the new name +
         contact line, preserving the originals' run formatting.
      4. For each known section header found in the source, blow away its
         body paragraphs and re-emit content using cloned styles.
    """
    doc = Document(io.BytesIO(source_bytes))
    if doc.tables:
        raise ValueError("source docx uses tables — cannot edit in place safely")

    paras = list(doc.paragraphs)
    if not paras:
        raise ValueError("source docx has no paragraphs")

    # ── header: name + contact ──────────────────────────────────────────────
    head_paras = [p for p in paras if p.text.strip()][:2]
    if sections.get("name") and len(head_paras) >= 1:
        _set_paragraph_text(head_paras[0], str(sections["name"]).strip())
    if sections.get("contact") and len(head_paras) >= 2:
        _set_paragraph_text(head_paras[1], str(sections["contact"]).strip())

    # ── body sections ───────────────────────────────────────────────────────
    section_ranges = _identify_sections(paras)
    if not section_ranges:
        raise ValueError("could not identify any section headers in source docx")

    for canonical, (start_i, end_i) in section_ranges.items():
        data = sections.get(canonical)
        # Skills section: data is a list of category dicts; empty list -> skip
        if data is None or (isinstance(data, (list, str)) and not data):
            continue

        header_para = paras[start_i]
        body_paras = paras[start_i + 1:end_i]

        bullet_t, subhead_t, plain_t, fb = _capture_templates(body_paras)

        # Detach existing body
        for bp in body_paras:
            parent = bp._element.getparent()
            if parent is not None:
                parent.remove(bp._element)

        # Insert new content
        emitter = SECTION_EMITTERS.get(canonical)
        if emitter is None:
            continue
        cursor = header_para._element
        emitter(cursor, data, bullet_t, subhead_t, plain_t, fb)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def sections_to_text(sections: dict) -> str:
    """Render the structured sections dict to the plain-text markdown format
    consumed by the legacy `_build_docx` renderer and the chat preview.
    """
    out: list[str] = []
    if sections.get("name"):
        out.append(str(sections["name"]).strip())
    if sections.get("contact"):
        out.append(str(sections["contact"]).strip())
    out.append("")

    if sections.get("summary"):
        out.append("## SUMMARY")
        out.append(str(sections["summary"]).strip())
        out.append("")

    if sections.get("skills"):
        out.append("## SKILLS")
        for cat in sections["skills"]:
            cat_name = str(cat.get("category", "") or "").strip()
            items = [str(x).strip() for x in (cat.get("items", []) or []) if str(x).strip()]
            if not items:
                continue
            line = f"**{cat_name}:** {', '.join(items)}" if cat_name else ", ".join(items)
            out.append(line)
        out.append("")

    if sections.get("experience"):
        out.append("## EXPERIENCE")
        for item in sections["experience"]:
            head = _format_exp_head(item)
            if head:
                out.append(f"### {head}")
            for b in item.get("bullets", []) or []:
                text = str(b).strip()
                if text:
                    out.append(f"- {text}")
            out.append("")

    if sections.get("projects"):
        out.append("## PROJECTS")
        for item in sections["projects"]:
            head = _format_project_head(item)
            if head:
                out.append(f"### {head}")
            for b in item.get("bullets", []) or []:
                text = str(b).strip()
                if text:
                    out.append(f"- {text}")
            out.append("")

    if sections.get("education"):
        out.append("## EDUCATION")
        for item in sections["education"]:
            inst = str(item.get("institution", "") or "").strip()
            loc = str(item.get("location", "") or "").strip()
            head = f"{inst} — {loc}" if inst and loc else inst or loc
            if head:
                out.append(f"### {head}")
            body_bits = []
            deg = str(item.get("degree", "") or "").strip()
            dates = str(item.get("dates", "") or "").strip()
            if deg:
                body_bits.append(deg)
            if dates:
                body_bits.append(dates)
            if body_bits:
                out.append(" · ".join(body_bits))
        out.append("")

    if sections.get("certifications"):
        out.append("## CERTIFICATIONS")
        for item in sections["certifications"]:
            text = str(item).strip()
            if text:
                out.append(f"- {text}")
        out.append("")

    if sections.get("achievements"):
        out.append("## ACHIEVEMENTS")
        for item in sections["achievements"]:
            text = str(item).strip()
            if text:
                out.append(f"- {text}")
        out.append("")

    if sections.get("values_alignment"):
        out.append("## VALUES ALIGNMENT")
        for item in sections["values_alignment"]:
            text = str(item).strip()
            if text:
                out.append(f"- {text}")
        out.append("")

    if sections.get("additional_information"):
        out.append("## ADDITIONAL INFORMATION")
        text = str(sections["additional_information"]).strip()
        if text:
            out.append(text)
        out.append("")

    return "\n".join(out).rstrip() + "\n"
