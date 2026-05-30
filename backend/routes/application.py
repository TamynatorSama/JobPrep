"""
Application-tailoring route.

POST /application/tailor-resume — given a job (company, role, JD) and the user's
master resumes, ask Gemini 2.5 Pro to:
  1. pick the closest-matching master resume
  2. produce a tailored, ATS-friendly resume (returned as a .docx)
  3. produce a cover letter (streamed as tokens for the chat bubble)
  4. produce a scorecard (verbatim-match %, role-title alignment, etc.)

The handler is SSE and emits these event types:
  stage          banner (rendered into the logs sidebar by the Rust client)
  token          cover-letter prose (rendered into the chat bubble)
  scorecard      structured JSON (Rust renders as a card)
  resume_docx    base64-encoded .docx bytes (Rust saves to disk)
  done | error
"""
import asyncio
import base64
import io
import json
import re

from fastapi import APIRouter
from sse_starlette.sse import EventSourceResponse

from google import genai
from google.genai import types as genai_types
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from models import ApplicationRequest, KnockoutRequest
from routes.docx_editor import sections_to_text


router = APIRouter()

# -- Master prompt (verbatim from the user's spec) ---------------------------
MASTER_PROMPT = """Persona and Role:


* Act as a dual-expert consultant embodying the roles of a Senior Technical Recruiter, an ATS Logic Engineer, and a Strategic Career Hacker.


* Your objective is to engineer an application that mathematically maximizes the probability of an interview by exploiting the mechanics of Applicant Tracking Systems (ATS) like Greenhouse, Lever, and Workable, and the psychology of recruiters.


* You operate based on 'Verbatim Match', 'Hire/ No Hire', and 'Structure Mirroring' methodologies, treating a resume as a dedicated proof-of-match document for a specific Job Description (JD).




Operational Workflow:




PHASE 1: Semantic Extraction & Strategy Setup


1. Request Input: One or more master resumes are provided below. They belong to the SAME candidate -- treat them as ONE unified evidence pool, NOT as alternatives to pick one of. Aggregate the STRONGEST evidence per section from across ALL of them. The JD is provided too.


2. The Verbatim Audit: Scan the JD for specific technical phrases. Identify exact strings (e.g., 'manipulate large datasets' instead of 'big data handling').


3. Structure Analysis: Identify headers in the JD responsibilities (e.g., 'Data Strategy') to use as CV sub-headers later.


4. Values Detection: Flag terms from the JD's 'Values' or 'Culture' section for alignment.




PHASE 2: System-Beater Resume Optimization


Generate a revised, text-only resume using these rules:


* Benchmark: Infer the role family from the JD title (engineering, data, product, design, marketing, ops, finance, etc.) and use that family's strongest resume patterns as the benchmark. Do NOT assume a tech role.


* Rule 1 (Role Title Trick): Match the candidate's Resume Job Title to the JD title exactly, provided responsibilities align. Do not falsely upgrade seniority.


* Rule 2 (Structure Mirroring): Categorize bullet points under sub-headers that mirror the JD's section headers.


* Rule 3 (Verbatim Keyword Injection): Integrate identified exact phrases naturally into bullet points.


* Rule 4 (Values Alignment): Create a section titled 'Values Alignment' linking traits to company keywords.


* Rule 5 (Quantification): Apply the Google XYZ formula (Accomplished [X] as measured by [Y], by doing [Z]). Aim for a number (%, $, time) on every impact bullet. If the master resumes don't supply a real figure, you MAY add a conservative estimate — but MARK it approximate (e.g. "~30%", "(est.)"). NEVER present an invented figure as a precisely measured result.


* Rule 6 (Proactive Logistics): Add an 'Additional Information' section addressing location and visa status.


* Formatting: Single-column, ATS-friendly. No tables, or columns. No citation markers like '[cite start]'.




PHASE 3: Hook & Match Cover Letter


Draft a letter using this 5-paragraph formula:


1. The Hook: Open with enthusiasm for the specific company mission/product.


2. Experience Story: Provide a condensed, quantified STAR format example.


3. Why This Role: Mention specific excitement about JD responsibilities.


4. Close: Confirm location/availability and request a discussion.


* Tone: Human, natural, using contractions. Avoid 'I am writing to express my interest.'




PHASE 4: Final Output & Analysis


1. Output: Provide rewritten Resume and Cover Letter.


2. Scorecard: Technical Match: List skills using JD terminology. Provide an objective 'Hire' or 'No Hire' recommendation based on fit.


   'Verbatim Match Score' (0-100%),


   'Role Title Alignment' (Yes/No), and


   'Quantification Check'.


3. Next Step: Ask if the user wants to simulate a 'Knockout Question' screening.




Operational Guardrails:


* Honest embellishment (NOT fabrication): You MAY (a) surface skills clearly IMPLIED by the candidate's real experience even when not named verbatim, (b) reframe and strengthen existing bullets in the JD's wording, (c) write a compelling tailored summary and Values Alignment, and (d) add conservative, clearly-marked metric estimates per Rule 5.


* HARD LIMITS — never invent any of: employers, job titles, companies, degrees, certifications, dates, or projects/experiences that did not happen. If a required skill has NO basis in the candidate's real background, OMIT it — do not claim it. Strong but defensible in an interview is the bar.


* Formatting: .docx compatible. Do not include metadata or citation text.


* Tone: Professional, authoritative, and tactical coach.
"""

# Output schema the LLM must follow. Embedded into the user prompt because
# google-genai's `response_schema` rejects long descriptions.
OUTPUT_SHAPE = """\
Return ONLY a single JSON object (no markdown fences, no leading prose) with \
this exact shape:
{
  "resumes_used": ["<labels of every master resume that contributed at least one bullet/skill/fact>"],
  "aggregation_notes": "<1-3 sentences describing what came from where. Example: 'Experience from Resume A (recent SWE roles); Projects from Resume B (ML/AI shipped work); Education from Resume A.'>",
  "sections": {
    "name": "<candidate full name>",
    "contact": "<single line: email | phone | city, country | linkedin | github (omit any field that is not in the master resumes)>",
    "summary": "<2-3 line prose summary tailored to the JD role title. Mirror JD wording. No first-person 'I'. No filler adjectives. No contact info, no URLs, no 'Seeking...' opener.>",
    "skills": [
      { "category": "<category that fits THIS role>", "items": ["...", "..."] }
    ],
    "experience": [
      {
        "role":     "Senior Software Engineer",
        "company":  "Acme",
        "location": "Remote",
        "start":    "Jan 2023",
        "end":      "Present",
        "bullets":  ["Action verb + what + technology + outcome.", "..."]
      }
    ],
    "projects": [
      {
        "name":     "ProjectX",
        "subtitle": "RAG pipeline over Postgres",
        "date":     "Mar 2024",
        "bullets":  ["Built ... using ... to ...", "Tech: Python, FastAPI, pgvector"]
      }
    ],
    "education": [
      {
        "institution": "MIT",
        "location":    "Cambridge, MA",
        "degree":      "BS Computer Science",
        "dates":       "Aug 2019 – May 2023"
      }
    ],
    "certifications": ["AWS Solutions Architect Associate — Amazon · Jan 2024"],
    "achievements":   ["Achievement description · Year"],
    "values_alignment": ["Trait — mapped to company keyword/value", "..."],
    "additional_information": "<one short line: location, work auth / visa status, availability. Omit if nothing concrete from the master resumes or request payload supports it.>"
  },
  "cover_letter": "<full cover letter, 180-260 words, plain prose with paragraph breaks. Do NOT start with 'I am writing to express my interest'.>",
  "scorecard": {
    "verbatim_match_score": 0-100,
    "role_title_alignment": "Yes" | "No",
    "quantification_check": "Pass" | "Needs Work",
    "hire_recommendation":  "Hire" | "No Hire",
    "skills_matched":       ["skill1", "skill2", "..."]   // SHORT keyword skills only: 1-3 words each, max 25 chars. 6-10 skills total. NEVER full sentences.
  }
}

IMPORTANT:
  - Omit ANY section key entirely (do not emit empty arrays/strings) if the aggregated evidence pool has nothing for it.
  - Do NOT emit a top-level `resume_text` field — content is structured-only and the backend renders the .docx itself.
  - For each experience entry, prefer 3-4 bullets. For each project, 2-3 bullets max.
  - Mirror EXACT JD keyword spelling inside skills/bullets where the candidate actually possesses the skill.
  - SKILLS categories are NOT fixed — choose 4-6 categories that fit THIS role family. Engineering: Languages, Frameworks, Databases, Tools, AI/ML. Product: Product & Strategy, Analytics, Research, Tools. Marketing: Channels, Martech, Analytics, Content. Pick what matches; never force engineering buckets onto a non-engineering role.
  - Any estimated/inferred number MUST be marked approximate ("~", "(est.)"). Never fabricate a precise measured figure.
"""


@router.post("/tailor-resume")
async def tailor_resume(req: ApplicationRequest):
    async def generate():
        if not req.api_key:
            yield _evt("error", "No Gemini API key set. Open Settings → API Keys.")
            return
        if not req.master_resumes:
            yield _evt("error", "No master resumes provided. Add one in Settings → Resume.")
            return
        if not req.company.strip() or not req.role.strip():
            yield _evt("error", "Company and role are required.")
            return

        yield _evt("stage", "\n\n---\n📋 **Analyzing JD and picking the best master resume…**\n\n")
        await asyncio.sleep(0)

        prompt = _build_prompt(req)

        try:
            client = genai.Client(api_key=req.api_key)

            yield _evt("stage", "\n📝 **Drafting tailored resume + cover letter…**\n\n")
            await asyncio.sleep(0)

            # ── Pass 1: draft. Structured JSON — we don't token-stream this
            #    because mid-stream JSON is unparseable.
            try:
                draft_resp, model_name = await _generate_json(client, prompt, temperature=0.4)
            except Exception as exc:
                yield _evt("error", f"Model call failed: {exc}")
                return
            yield _evt("stage", f"\n🧠 **Model:** `{model_name}`\n\n")
            await asyncio.sleep(0)

            try:
                data = _loads_lenient(_response_text(draft_resp))
            except (json.JSONDecodeError, ValueError) as exc:
                yield _evt("error", f"Model returned non-JSON output: {exc}")
                return

            # ── Pass 2: self-critique + refine against the JD. Best-effort —
            #    a failed or garbled refine falls back to the draft so we never
            #    lose a usable result.
            yield _evt("stage", "\n🔬 **Refining against the JD (closing keyword gaps + tightening bullets)…**\n\n")
            await asyncio.sleep(0)
            try:
                refine_prompt = _build_refine_prompt(req, json.dumps(data, ensure_ascii=False))
                refine_resp, _ = await _generate_json(client, refine_prompt, temperature=0.3)
                refined = _loads_lenient(_response_text(refine_resp))
                if isinstance(refined, dict) and refined.get("sections"):
                    data = refined
                    yield _evt("stage", "\n✅ **Refinement applied.**\n\n")
                else:
                    yield _evt("stage", "\n↩️ Refinement skipped (no sections in output) — kept the draft.\n\n")
            except Exception as exc:
                print(f"[application] refine pass failed: {exc!r}", flush=True)
                yield _evt("stage", f"\n↩️ Refinement skipped ({type(exc).__name__}) — kept the draft.\n\n")
            await asyncio.sleep(0)

            sections     = data.get("sections") or {}
            cover_letter = (data.get("cover_letter") or "").strip()
            scorecard    = data.get("scorecard")    or {}
            resumes_used = data.get("resumes_used") or []
            agg_notes    = (data.get("aggregation_notes") or "").strip()

            # Backwards-compat with older model outputs.
            if not resumes_used and data.get("selected_resume_label"):
                resumes_used = [data["selected_resume_label"]]

            # Build plain-text view from sections for chat preview / knockout.
            resume_text = sections_to_text(sections).strip() if sections else ""

            # Legacy fallback: if a model still emits resume_text directly.
            if not resume_text and data.get("resume_text"):
                resume_text = str(data["resume_text"]).strip()

            if not resume_text or not cover_letter:
                yield _evt("error", "Model output was missing the resume sections or cover letter.")
                return

            used_label = ", ".join(resumes_used) if resumes_used else "(unspecified)"
            yield _evt("stage", f"\n✨ **Aggregated from:** _{used_label}_\n\n")
            if agg_notes:
                yield _evt("stage", f"_{agg_notes}_\n\n")
            await asyncio.sleep(0)

            # ── Stream the cover letter as tokens (word-by-word) so the chat
            #    bubble fills in incrementally — same UX as the chat route.
            yield _evt("stage", "\n💌 **Cover Letter**\n\n")
            for word in re.split(r"(\s+)", cover_letter):
                if word:
                    yield _evt("token", word)
                    await asyncio.sleep(0.005)
            yield _evt("token", "\n\n")

            # ── Build the .docx and send it back as base64 ──────────────────
            #    Always render from scratch via `_build_docx`. The in-place
            #    "duplicate base resume" path was dropped because it silently
            #    skipped sections whose headers weren't in the source.
            yield _evt("stage", "\n📄 **Building tailored resume.docx…**\n\n")
            await asyncio.sleep(0)
            docx_bytes = _build_docx(resume_text)
            yield _evt(
                "resume_docx",
                base64.b64encode(docx_bytes).decode("ascii"),
            )

            # ── Scorecard ───────────────────────────────────────────────────
            # Bundle resumes_used + aggregation_notes into the scorecard so the
            # frontend's persistent `job.scorecard` carries them. Stage banners
            # show the same info but disappear from view once streaming ends.
            yield {"data": json.dumps({"type": "scorecard", "content": {
                **scorecard,
                "resumes_used": resumes_used,
                "aggregation_notes": agg_notes,
            }})}
            await asyncio.sleep(0)

            # ── Pass the tailored resume text back so the client can hand
            #    it to the company-research step (used as background context).
            yield {"data": json.dumps({"type": "tailored_resume", "content": resume_text})}
            await asyncio.sleep(0)

            yield _evt("done", "")

        except Exception as exc:
            yield _evt("error", f"Application tailor error: {exc}")

    return EventSourceResponse(generate())


# ── Knockout-question screening ──────────────────────────────────────────────

KNOCKOUT_PROMPT = """\
You are a senior technical recruiter at the target company running a 15–20
minute KNOCKOUT phone screen. The explicit purpose of this call is to
DISQUALIFY weak candidates before they waste the hiring manager's time.
You've read the JD top-to-bottom and the candidate's tailored resume. You
are tactical, time-pressured, and you do not tolerate fluff.

# What to produce

Predict the EXACT 6–8 questions you would actually ask in this screen, in
the order you'd ask them. Your output is a written prep dossier the
candidate will study before the real call.

# Question mix (this is the spine — pick from these buckets so the screen feels real, not a checklist)

(A) LOGISTICS — 1–2 questions. Location/relocation, work authorization /
visa, salary expectations vs the JD's posted band (or current market for
this title), notice period, remote/hybrid commitment if specified in the JD.

(B) MUST-HAVE TECHNICAL — 2–3 questions. Identify the 2–3 hardest technical
requirements in the JD. For each, write a question whose answer immediately
reveals whether the candidate ACTUALLY has the skill versus just listed the
keyword.
  - For "experience with distributed systems": "Walk me through the last
    time you debugged a distributed-system issue in production. What was
    the failure mode and how did you identify it?"
  - For "SQL proficiency": "Describe a query you wrote in the last month
    that was painful to optimize. What was the bottleneck?"
  - For "ML model deployment": "Tell me about a model you actually deployed.
    What monitoring did you set up, and what alerted you to model drift?"
  ABSOLUTELY DO NOT ask textbook trivia ("what is normalization?",
  "explain TCP/IP"). Every technical question MUST ask for a real story or
  a concrete past experience.

(C) BEHAVIORAL RED-FLAG CHECK — 1 question. Scan the resume for the actual
red flag (short tenure, employment gap, lateral move, industry switch, no
recent achievements, suspicious title progression). Ask the question that
surfaces whether it's a real problem.
  Example: "I see you were at [X] for 8 months between [Y] and [Z] — walk
  me through what happened."

(D) SCENARIO / ROLE-PLAY — 1 question. Give a realistic on-the-job
scenario the candidate would hit in this exact role and ask how they'd
handle it. The scenario MUST be specific to what THIS company does, not
generic. Use the company's domain.
  Example for a fintech SWE: "Our primary payment processor returns 500s
  during peak Friday volume. You're the on-call engineer. Walk me through
  your first 10 minutes."
  Example for an ML engineer at a recsys company: "Your A/B test shows the
  new ranking model has 4% higher engagement but 12% more 'show fewer like
  this' clicks. What do you do?"

(E) MOTIVATION & CULTURE FIT — 1 question. Why THIS company specifically
(not "I want to grow my skills"), grounded in something concrete: a product,
recent announcement, mission element. One culture-fit probe if the company
has known cultural values worth testing.

# Output format

For each question, output EXACTLY this structure:

## Question N — [TYPE in caps, e.g. TECHNICAL or SCENARIO]

> [The actual question verbatim, as if a recruiter is speaking. One question. Conversational tone.]

*Why they're asking:* [One sentence — the specific red flag this surfaces.]

**Your answer:** [Concrete answer drawn from the resume. Name the specific project, the specific tech, the specific metric. 3–6 sentences. If the resume doesn't fully support an answer, lead with "**GAP — **" and write an honest one-sentence frame that does not lie.]

**Likely follow-up:** [The probing question the recruiter will ask next based on your answer. Plus a one-sentence handle for it.]

**What NOT to say:** [The vague / buzzword / generic answer that would fail this question.]

After the questions, output a final section:

# Verdict

**Predicted outcome:** Pass / Borderline / Fail
**The one weakness:** [The single biggest hole between this resume and this JD. One sentence, brutally honest.]
**Tonight's prep:** [The ONE thing to rehearse before the actual call. Specific, actionable, under 30 minutes of work.]

# Absolute rules

- Never invent resume content. Every "Your answer" must cite something that's actually in the resume.
- Never use stock interview questions ("tell me about yourself", "what's your greatest weakness", "where do you see yourself in 5 years"). The candidate has heard them all.
- Every question must be specific to THIS candidate + THIS role + THIS company.
- The order matters — open with logistics, build through technical, end on motivation.
- Single questions only. No multi-part. No bullet-list questions.
- Stream the output as markdown — H2 per question, H1 only for the final Verdict section.
"""


@router.post("/knockout-screen")
async def knockout_screen(req: KnockoutRequest):
    async def generate():
        if not req.api_key:
            yield _evt("error", "No Gemini API key set. Open Settings → API Keys.")
            return
        if not req.tailored_resume.strip():
            yield _evt(
                "error",
                "Run the Application Prep step first — the knockout screen "
                "needs the tailored resume as input.",
            )
            return
        if not req.company.strip() or not req.role.strip():
            yield _evt("error", "Company and role are required.")
            return

        yield _evt("stage", "\n\n---\n🎯 **Simulating recruiter knockout screen…**\n\n")
        await asyncio.sleep(0)

        loc = f"\nLocation: {req.location}" if req.location else ""
        contents = (
            KNOCKOUT_PROMPT
            + "\n\n=== JOB ===\n"
            + f"Company: {req.company}\nRole: {req.role}{loc}\n\n"
            + f"=== JOB DESCRIPTION ===\n{req.job_description}\n\n"
            + f"=== CANDIDATE'S TAILORED RESUME ===\n{req.tailored_resume}\n"
        )

        try:
            client = genai.Client(api_key=req.api_key)

            # Try Gemini 3 first for sharper recruiter judgement; fall back
            # gracefully when a model name isn't available on the account.
            candidate_models = [
                "gemini-3.1-pro-preview",
                "gemini-3-pro-preview",
                "gemini-2.5-pro",
            ]
            stream = None
            last_err: Exception | None = None
            for model_name in candidate_models:
                try:
                    stream = await client.aio.models.generate_content_stream(
                        model=model_name,
                        contents=contents,
                        config=genai_types.GenerateContentConfig(temperature=0.5),
                    )
                    yield _evt("stage", f"\n🧠 **Model:** `{model_name}`\n\n")
                    await asyncio.sleep(0)
                    break
                except Exception as exc:
                    last_err = exc
                    yield _evt("stage", f"\n⚠️ {model_name} unavailable — trying next…\n\n")
                    await asyncio.sleep(0)
            if stream is None:
                raise last_err or RuntimeError("No model available")

            # True streaming — these questions form a narrative the user wants
            # to watch render in the chat bubble.
            async for chunk in stream:
                text = _response_text(chunk)
                if text:
                    yield _evt("token", text)
                    await asyncio.sleep(0)

            yield _evt("done", "")

        except Exception as exc:
            yield _evt("error", f"Knockout screen error: {exc}")

    return EventSourceResponse(generate())


# ── helpers ──────────────────────────────────────────────────────────────────

def _evt(etype: str, content: str) -> dict:
    return {"data": json.dumps({"type": etype, "content": content})}


def _strip_fences(text: str) -> str:
    t = text.strip()
    t = re.sub(r"^```(?:json)?\s*", "", t, flags=re.MULTILINE)
    t = re.sub(r"\s*```$", "", t, flags=re.MULTILINE)
    return t.strip()


def _loads_lenient(text: str) -> dict:
    """Parse the first complete JSON object out of a model response.

    Tolerates leading prose, code fences, and — importantly — trailing junk
    after the object (Gemini sometimes appends stray `}` braces, which makes
    a plain `json.loads` raise "Extra data"). Uses `raw_decode` so we take the
    first balanced object and ignore whatever follows.
    """
    t = _strip_fences(text)
    start = t.find("{")
    if start == -1:
        raise ValueError("no JSON object found in model output")
    obj, _ = json.JSONDecoder().raw_decode(t[start:])
    return obj


def _response_text(response) -> str:
    if hasattr(response, "text") and response.text:
        return response.text
    try:
        parts = response.candidates[0].content.parts
        return "".join(getattr(p, "text", "") for p in parts)
    except Exception:
        return str(response)


def _build_prompt(req: ApplicationRequest) -> str:
    def _format_resume(i: int, r) -> str:
        return (
            f"=== MASTER RESUME [{i+1}/{len(req.master_resumes)}] "
            f"LABEL: {r.name} ===\n{r.text}"
        )

    resumes_block = "\n\n".join(
        _format_resume(i, r) for i, r in enumerate(req.master_resumes)
    )
    loc = f"\nLocation: {req.location}" if req.location else ""
    aggregation_note = (
        f"\n\nAGGREGATION REMINDER: {len(req.master_resumes)} master resumes "
        "below belong to the SAME candidate. Merge them into ONE unified "
        "evidence pool per PHASE 1. Pull the STRONGEST item per section from "
        "across ALL of them. Use every contributing resume's LABEL in "
        "`resumes_used`."
        if len(req.master_resumes) > 1
        else "\n\nNOTE: Only one master resume supplied. Use it as the sole "
             "evidence pool."
    )
    return (
        MASTER_PROMPT
        + "\n\n=== JOB CONTEXT ===\n"
        + f"Company: {req.company}\n"
        + f"Role: {req.role}{loc}\n\n"
        + f"=== JOB DESCRIPTION ===\n{req.job_description}\n\n"
        + f"=== AVAILABLE MASTER RESUMES ({len(req.master_resumes)}) ===\n"
        + resumes_block
        + aggregation_note
        + "\n\n=== TASK ===\n"
        + OUTPUT_SHAPE
    )


REFINE_PROMPT = """You are a Senior Technical Recruiter + ATS Logic Engineer doing a \
SECOND-PASS review of a tailored-resume DRAFT (JSON below) against the target JD. Make it \
materially stronger WITHOUT fabricating. Critique the draft, then return an IMPROVED version \
applying every fix:

1. Keyword coverage: scan the JD for must-have phrases; every one the candidate plausibly \
satisfies must appear VERBATIM in skills or a bullet. Close the gaps.
2. Bullets: each experience/project bullet follows XYZ (Accomplished X, measured by Y, by \
doing Z) and carries a number. If no real figure exists, add a conservative estimate MARKED \
approximate ("~", "(est.)") — never a fake precise metric. Cut generic or duplicate bullets.
3. Summary: 2-3 lines, mirrors the JD title + its top 3 requirements, no first-person, no filler.
4. Skills: categories must fit THIS role family (don't force engineering buckets on a \
non-engineering role); 4-6 categories; JD wording where the candidate has the skill.
5. Values Alignment: map REAL candidate traits to the company's stated values/keywords.
6. Cover letter: keep the 5-beat structure + human tone; no "I am writing to express my interest".

HARD LIMITS (same as the draft): never invent employers, titles, companies, degrees, \
certifications, dates, or projects. Preserve `resumes_used`; update `aggregation_notes` only if \
your edits shift what came from where.

Return ONLY the improved JSON in the exact shape specified below.
"""


def _build_refine_prompt(req: ApplicationRequest, draft_json: str) -> str:
    loc = f"\nLocation: {req.location}" if req.location else ""
    return (
        REFINE_PROMPT
        + "\n\n=== JOB CONTEXT ===\n"
        + f"Company: {req.company}\nRole: {req.role}{loc}\n\n"
        + f"=== JOB DESCRIPTION ===\n{req.job_description}\n\n"
        + "=== DRAFT TO IMPROVE (JSON) ===\n"
        + draft_json
        + "\n\n=== REQUIRED OUTPUT SHAPE ===\n"
        + OUTPUT_SHAPE
    )


async def _generate_json(client, contents: str, *, temperature: float):
    """Call Gemini for JSON output, trying preview models then stable so a
    404 / model-not-found doesn't block the feature. Returns
    (response, model_name); raises the last error if every model fails."""
    candidate_models = [
        "gemini-3.1-pro-preview",
        "gemini-3-pro-preview",
        "gemini-2.5-pro",
    ]
    last_err: Exception | None = None
    for model_name in candidate_models:
        try:
            resp = await client.aio.models.generate_content(
                model=model_name,
                contents=contents,
                config=genai_types.GenerateContentConfig(
                    response_mime_type="application/json",
                    temperature=temperature,
                ),
            )
            return resp, model_name
        except Exception as exc:
            last_err = exc
    raise last_err or RuntimeError("No model available")


def _add_bottom_border(paragraph) -> None:
    """Attach a full-width horizontal rule under a paragraph (section divider).

    `w:sz` is in eighths-of-a-point — 8 = 1pt. `w:space` (in points) pushes
    the rule a few pt below the text so the heading doesn't sit on the line.
    """
    pPr = paragraph._p.get_or_add_pPr()
    # Drop any existing pBdr so we don't stack borders.
    for existing in pPr.findall(qn("w:pBdr")):
        pPr.remove(existing)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "8")        # 1pt — visible without screaming
    bottom.set(qn("w:space"), "4")        # 4pt gap between text and rule
    bottom.set(qn("w:color"), "606060")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_run_with_inline_bold(paragraph, text: str) -> None:
    """Add runs to `paragraph` splitting `text` on '**...**' bold spans."""
    parts = text.split("**")
    for i, seg in enumerate(parts):
        if not seg:
            continue
        r = paragraph.add_run(seg)
        r.bold = (i % 2 == 1)


FONT_NAME = "Times New Roman"


def _set_font(run, *, size_pt: float | None = None, bold: bool | None = None,
              color: "RGBColor | None" = None) -> None:
    """Apply Times New Roman + optional size/bold/color to a run.

    python-docx does not always propagate `font.name` from the Normal style
    into runs with explicit rPr children, so we set it on every run we
    build to guarantee the whole resume renders in one font.
    """
    run.font.name = FONT_NAME
    # rFonts override — pins eastAsia + cs slots so Word doesn't substitute
    # back to Calibri for any character class.
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _build_docx(resume_text: str) -> bytes:
    """Render the LLM's plain-text resume into a clean, ATS-friendly .docx.

    Layout rules:
        Line 1            → candidate name (centered, 20pt, bold)
        Line 2            → contact info (centered, 10pt, grey)
        '## SECTION'      → uppercase heading, 11.5pt bold, bottom-rule
        '### Subhead'     → 11pt bold (role / project / school)
        '- bullet'        → List Bullet style with hanging indent
        '**text**'        → inline bold
        '' (blank)        → tight spacer

    Times New Roman 11pt + 0.6in margins is the most ATS-portable combo.
    """
    doc = Document()

    # Tighter margins: 0.6in left/right, 0.5in top/bottom.
    section = doc.sections[0]
    section.top_margin    = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin   = Inches(0.6)
    section.right_margin  = Inches(0.6)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after  = Pt(2)
    normal.paragraph_format.line_spacing = 1.15

    lines = [ln.rstrip() for ln in resume_text.split("\n")]
    saw_name = False
    saw_contact = False

    GREY = RGBColor(0x55, 0x55, 0x55)
    BLACK_INK = RGBColor(0x10, 0x10, 0x10)

    for raw in lines:
        line = raw.strip()

        if not line:
            spacer = doc.add_paragraph("")
            spacer.paragraph_format.space_after = Pt(2)
            continue

        # ── Header: name + contact (first two non-empty lines) ───────────
        if not saw_name:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
            _set_font(p.add_run(line), size_pt=20, bold=True, color=BLACK_INK)
            saw_name = True
            continue

        if not saw_contact:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(4)
            p.alignment = 1
            _set_font(p.add_run(line), size_pt=10, color=GREY)
            saw_contact = True
            continue

        # ── Section heading (## or #) ────────────────────────────────────
        if line.startswith("## "):
            p = doc.add_paragraph()
            # Modest space above so sections feel separated without burning
            # vertical real estate; smaller gap below so the rule + body
            # read as one block.
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(3)
            _set_font(p.add_run(line[3:].strip().upper()),
                      size_pt=11.5, bold=True, color=BLACK_INK)
            _add_bottom_border(p)
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            _set_font(p.add_run(line[2:].strip()), size_pt=13, bold=True)
            continue

        # ── Sub-heading (### role/project/school) ─────────────────────────
        if line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(1)
            _add_run_with_inline_bold(p, line[4:].strip())
            # Force the whole sub-head bold regardless of inline markers.
            for r in p.runs:
                _set_font(r, size_pt=11, bold=True)
            continue

        # ── Bullet ────────────────────────────────────────────────────────
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            _add_run_with_inline_bold(p, line[2:].strip())
            for r in p.runs:
                _set_font(r)
            continue

        # ── Bare bold line e.g. "**Education**" ───────────────────────────
        if line.startswith("**") and line.endswith("**") and len(line) > 4:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(1)
            _set_font(p.add_run(line[2:-2]), bold=True)
            continue

        # ── Regular paragraph (with optional inline **bold** spans) ──────
        p = doc.add_paragraph()
        _add_run_with_inline_bold(p, line)
        for r in p.runs:
            _set_font(r)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
